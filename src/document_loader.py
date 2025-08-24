"""
Document loader module for DocuChat RAG system.
Handles file discovery and text extraction from various document formats.
"""

import os
import sys
import warnings
from pathlib import Path
from typing import List, Dict, Generator, Optional, Tuple
import traceback

# Import document processing libraries
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None


class DocumentLoader:
    """Handles document loading and text extraction from various file formats."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.md', '.docx'}
    
    def __init__(self, verbose: bool = False):
        """
        Initialize the document loader.
        
        Args:
            verbose: Enable verbose output
        """
        self.verbose = verbose
        self._check_dependencies()
    
    def _check_dependencies(self) -> None:
        """Check if required dependencies are available."""
        missing_deps = []
        
        if pypdf is None:
            missing_deps.append("pypdf==6.0.0")
        
        if docx is None:
            missing_deps.append("python-docx==1.2.0")
        
        if missing_deps:
            raise ImportError(
                f"Missing required dependencies: {', '.join(missing_deps)}. "
                f"Install with: pip install {' '.join(missing_deps)}"
            )
    
    def _log(self, message: str) -> None:
        """Log message if verbose mode is enabled."""
        if self.verbose:
            print(f"[DocumentLoader] {message}")
    
    def discover_files(self, directory: str) -> List[Path]:
        """
        Recursively discover all supported document files in a directory.
        
        Args:
            directory: Root directory to search
            
        Returns:
            List of Path objects for supported files
        """
        directory_path = Path(directory)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        if not directory_path.is_dir():
            raise NotADirectoryError(f"Path is not a directory: {directory}")
        
        self._log(f"Discovering files in: {directory_path.resolve()}")
        
        discovered_files = []
        
        try:
            # Use rglob for recursive search
            for file_path in directory_path.rglob('*'):
                if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                    discovered_files.append(file_path)
                    self._log(f"Found: {file_path}")
        
        except PermissionError as e:
            self._log(f"Permission denied accessing directory: {e}")
            warnings.warn(f"Permission denied: {e}", UserWarning)
        
        except Exception as e:
            self._log(f"Error during file discovery: {e}")
            warnings.warn(f"File discovery error: {e}", UserWarning)
        
        # Sort files for consistent ordering
        discovered_files.sort(key=lambda p: str(p).lower())
        
        self._log(f"Discovered {len(discovered_files)} supported files")
        return discovered_files
    
    def load_text_file(self, file_path: Path) -> str:
        """
        Load text from a plain text file (.txt, .md).
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If file cannot be read
        """
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    self._log(f"Successfully read {file_path} with {encoding} encoding")
                    return content
            
            except UnicodeDecodeError:
                continue
            
            except Exception as e:
                raise Exception(f"Error reading text file {file_path}: {e}")
        
        raise Exception(f"Could not decode text file {file_path} with any encoding")
    
    def load_pdf_file(self, file_path: Path) -> str:
        """
        Load text from a PDF file using pypdf.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If PDF cannot be processed
        """
        if pypdf is None:
            raise ImportError("pypdf library not available")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                if pdf_reader.is_encrypted:
                    # Try to decrypt with empty password
                    if not pdf_reader.decrypt(""):
                        raise Exception(f"PDF is password protected: {file_path}")
                
                text_content = []
                total_pages = len(pdf_reader.pages)
                
                self._log(f"Processing PDF with {total_pages} pages: {file_path}")
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                        
                        if self.verbose and page_num % 10 == 0:
                            self._log(f"Processed {page_num}/{total_pages} pages")
                    
                    except Exception as e:
                        self._log(f"Error extracting text from page {page_num}: {e}")
                        continue
                
                extracted_text = '\n\n'.join(text_content)
                self._log(f"Successfully extracted {len(extracted_text)} characters from PDF")
                
                return extracted_text
        
        except Exception as e:
            raise Exception(f"Error processing PDF {file_path}: {e}")
    
    def load_docx_file(self, file_path: Path) -> str:
        """
        Load text from a DOCX file using python-docx.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If DOCX cannot be processed
        """
        if docx is None:
            raise ImportError("python-docx library not available")
        
        try:
            doc = docx.Document(file_path)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            extracted_text = '\n\n'.join(text_content)
            self._log(f"Successfully extracted {len(extracted_text)} characters from DOCX")
            
            return extracted_text
        
        except Exception as e:
            raise Exception(f"Error processing DOCX {file_path}: {e}")
    
    def load_document(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """
        Load text from a document file, automatically detecting the format.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (filename, extracted_text) or None if loading failed
        """
        if not file_path.exists():
            self._log(f"File not found: {file_path}")
            return None
        
        if not file_path.is_file():
            self._log(f"Path is not a file: {file_path}")
            return None
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.SUPPORTED_EXTENSIONS:
            self._log(f"Unsupported file type: {file_extension}")
            return None
        
        self._log(f"Loading document: {file_path}")
        
        try:
            if file_extension in {'.txt', '.md'}:
                text_content = self.load_text_file(file_path)
            elif file_extension == '.pdf':
                text_content = self.load_pdf_file(file_path)
            elif file_extension == '.docx':
                text_content = self.load_docx_file(file_path)
            else:
                raise Exception(f"Unsupported file extension: {file_extension}")
            
            if not text_content.strip():
                self._log(f"Warning: No text content extracted from {file_path}")
                return None
            
            self._log(f"Successfully loaded {len(text_content)} characters from {file_path}")
            return (str(file_path), text_content)
        
        except Exception as e:
            error_msg = f"Error loading {file_path}: {e}"
            self._log(error_msg)
            warnings.warn(error_msg, UserWarning)
            
            if self.verbose:
                traceback.print_exc()
            
            return None
    
    def load_documents_from_directory(self, directory: str) -> Generator[Tuple[str, str], None, None]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory: Directory path to process
            
        Yields:
            Tuples of (filename, extracted_text) for successfully loaded documents
        """
        discovered_files = self.discover_files(directory)
        
        if not discovered_files:
            self._log("No supported files found in directory")
            return
        
        total_files = len(discovered_files)
        loaded_count = 0
        error_count = 0
        
        self._log(f"Starting to process {total_files} files")
        
        for i, file_path in enumerate(discovered_files, 1):
            self._log(f"Processing file {i}/{total_files}: {file_path.name}")
            
            result = self.load_document(file_path)
            
            if result is not None:
                loaded_count += 1
                yield result
            else:
                error_count += 1
            
            if self.verbose and i % 5 == 0:
                self._log(f"Progress: {i}/{total_files} files processed")
        
        self._log(f"Processing complete: {loaded_count} loaded, {error_count} errors")


def main():
    """Simple test function for the document loader."""
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python document_loader.py <directory>")
        sys.exit(1)
    
    directory = sys.argv[1]
    loader = DocumentLoader(verbose=True)
    
    try:
        for filename, content in loader.load_documents_from_directory(directory):
            print(f"\n{'='*50}")
            print(f"File: {filename}")
            print(f"Content length: {len(content)} characters")
            print(f"Preview: {content[:200]}...")
    
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()